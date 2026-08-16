from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "smart_evacuation_route_planner_summary.md"
OUTPUT = ROOT / "docs" / "Smart_Evacuation_Route_Planner_Summary.docx"

FONT = "Calibri"
NAVY = RGBColor(22, 55, 84)
TEAL = RGBColor(0, 111, 115)
INK = RGBColor(35, 41, 47)
MUTED = RGBColor(91, 99, 108)
PALE = "EDF5F5"


def set_font(run, size, *, bold=False, italic=False, color=INK, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_columns(section, count=2, space_twips=360):
    section_properties = section._sectPr
    columns = section_properties.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        section_properties.append(columns)
    columns.set(qn("w:num"), str(count))
    columns.set(qn("w:space"), str(space_twips))


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, display, end])
    set_font(run, 8.5, color=MUTED)


def add_inline_markdown(paragraph, text, size=10.2):
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size - 0.3, color=TEAL, name="Consolas")
        else:
            run = paragraph.add_run(part)
            set_font(run, size)


def shade_paragraph(paragraph, fill):
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.66)
    section.bottom_margin = Inches(0.64)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.26)
    section.footer_distance = Inches(0.26)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.05

    heading = document.styles["Heading 3"]
    heading.font.name = FONT
    heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading.font.size = Pt(12)
    heading.font.bold = True
    heading.font.color.rgb = NAVY
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(3)
    heading.paragraph_format.keep_with_next = True


def add_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("Smart Evacuation Route Planner")
    set_font(left, 8.5, bold=True, color=NAVY)
    tab = paragraph.add_run("\tPlain-English Brief")
    set_font(tab, 8.5, color=MUTED)
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(7.05))

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    label = paragraph.add_run("Page ")
    set_font(label, 8.5, color=MUTED)
    add_field(paragraph, "PAGE")
    label = paragraph.add_run(" of ")
    set_font(label, 8.5, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def parse_blocks(markdown_text):
    lines = markdown_text.splitlines()
    title = lines[0].removeprefix("# ").strip()
    subtitle = lines[2].removeprefix("## ").strip()
    blocks = []
    paragraph = []
    for line in lines[4:]:
        stripped = line.strip()
        if not stripped:
            if paragraph:
                blocks.append(("paragraph", " ".join(paragraph)))
                paragraph = []
            continue
        if stripped.startswith("### "):
            if paragraph:
                blocks.append(("paragraph", " ".join(paragraph)))
                paragraph = []
            blocks.append(("heading", stripped[4:].strip()))
        else:
            paragraph.append(stripped)
    if paragraph:
        blocks.append(("paragraph", " ".join(paragraph)))
    return title, subtitle, blocks


def build():
    title, subtitle, blocks = parse_blocks(SOURCE.read_text(encoding="utf-8"))
    document = Document()
    document.core_properties.title = title
    document.core_properties.subject = "Plain-English summary of the Smart Evacuation Route Planner"
    document.core_properties.author = "Project Team"
    configure_styles(document)
    first_section = document.sections[0]
    configure_section(first_section)
    add_header_footer(first_section)

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(4)
    title_paragraph.paragraph_format.space_after = Pt(1)
    title_run = title_paragraph.add_run(title)
    set_font(title_run, 20.5, bold=True, color=NAVY)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.paragraph_format.space_after = Pt(6)
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    set_font(subtitle_run, 10.5, bold=True, color=TEAL)

    # Keep the opening explanation full width, then use a readable two-column brief.
    opening_heading = blocks.pop(0)
    opening_paragraphs = [blocks.pop(0), blocks.pop(0)]
    heading_paragraph = document.add_paragraph(style="Heading 3")
    heading_paragraph.add_run(opening_heading[1])
    for _, text in opening_paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.02
        add_inline_markdown(paragraph, text, 10.4)

    second_section = document.add_section(WD_SECTION.CONTINUOUS)
    configure_section(second_section)
    set_columns(second_section, 2, 300)
    second_section.header.is_linked_to_previous = True
    second_section.footer.is_linked_to_previous = True

    for kind, text in blocks:
        if kind == "heading":
            paragraph = document.add_paragraph(style="Heading 3")
            paragraph.add_run(text)
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        paragraph.paragraph_format.keep_together = False
        if text.startswith("`") and text.endswith("`"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.left_indent = Inches(0.08)
            paragraph.paragraph_format.right_indent = Inches(0.08)
            shade_paragraph(paragraph, PALE)
        add_inline_markdown(paragraph, text)

    # A trailing continuous section makes Word and LibreOffice balance the
    # final two columns instead of leaving the second column mostly empty.
    closing_section = document.add_section(WD_SECTION.CONTINUOUS)
    configure_section(closing_section)
    set_columns(closing_section, 1)
    closing_section.header.is_linked_to_previous = True
    closing_section.footer.is_linked_to_previous = True

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
