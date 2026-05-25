from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_font(run, name="Times New Roman", size=12, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def set_cell_width(cell, width):
    cell.width = Inches(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    tbl.insert(1, grid)


def repair_known_table_text(table):
    if not table.rows:
        return

    header = [cell.text.strip() for cell in table.rows[0].cells]
    if header == ["Crowd level", "Approximate density", "Route-planning meaning"]:
        values = [
            ("Low", "0.5 to 1.5 persons/m²", "Movement is possible with limited interference"),
            ("Medium", "1.5 to 3 persons/m²", "Movement is affected and route cost should increase"),
            ("High", "3 to 4 persons/m²", "Congestion risk is significant"),
            ("Critical", "Above 4 to 5 persons/m²", "Area should be avoided when alternatives exist"),
        ]
        for row, row_values in zip(table.rows[1:], values):
            for cell, value in zip(row.cells, row_values):
                cell.text = value

    if header == ["Step", "Process", "Description"]:
        for index, row in enumerate(table.rows[1:], start=1):
            row.cells[0].text = str(index)


def main():
    if len(sys.argv) != 2:
        raise SystemExit("Usage: polish_latex_docx.py <docx-path>")

    docx = Path(sys.argv[1])
    doc = Document(docx)
    style_names = {s.name for s in doc.styles}

    for section in doc.sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for style_name in ["Normal", "Body Text"]:
        if style_name in style_names:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(12)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.space_after = Pt(6)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in style_names:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(12)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)

    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            continue

        if idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(12)
            for run in p.runs:
                set_font(run, size=14, bold=True)
            continue

        if idx in (1, 2):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                set_font(run, size=12)
            continue

        if text.startswith("Table ") or text.startswith("Figure "):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                set_font(run, size=11)
            continue

        if "https://doi.org/" in text or text.startswith((
            "Chen,", "Gao,", "Haghani,", "Helbing,", "Fruin,", "Hevner,", "Kim,",
            "Lopez-", "Łukasik,", "Mocanu,", "Page,", "Peffers,", "Senanayake,",
            "Yiğit,", "Yin,", "Zhou,"
        )):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            set_font(run, size=12, bold=run.bold, italic=run.italic)

    for table in doc.tables:
        repair_known_table_text(table)
        if table.rows:
            mark_header(table.rows[0])
        if "Table Grid" in style_names:
            table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False
        col_count = len(table.columns)
        if col_count == 1:
            widths = [6.27]
        elif col_count == 2:
            widths = [1.9, 4.37]
        elif col_count == 3:
            widths = [1.25, 1.75, 3.27]
        elif col_count == 5:
            widths = [1.05, 1.35, 1.35, 1.35, 1.17]
        else:
            widths = [6.27 / col_count] * col_count
        set_table_grid(table, widths)
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                set_cell_width(cell, widths[min(c_idx, len(widths) - 1)])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        set_font(run, size=9, bold=(r_idx == 0 or run.bold))

    doc.save(docx)


if __name__ == "__main__":
    main()
